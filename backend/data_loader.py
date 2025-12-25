"""
Data Loader Module
Loads and processes both Excel and DOCX source files.
"""
import openpyxl
from docx import Document
from typing import List, Dict, Tuple, Any
import re


class DataLoader:
    """Handles loading and preprocessing of source documents."""
    
    def __init__(self, excel_path: str, docx_path: str):
        self.excel_path = excel_path
        self.docx_path = docx_path
        self.excel_data = [] # List of dicts
        self.doc_chunks = []
        
    def load_excel(self) -> List[Dict[str, Any]]:
        """Load Excel file via OpenPyXL and return as list of dicts."""
        wb = openpyxl.load_workbook(self.excel_path, data_only=True)
        ws = wb.active
        
        data = []
        rows = list(ws.iter_rows(values_only=True))
        
        if not rows:
            return []
            
        headers = [str(h).strip() for h in rows[0]]
        
        for row in rows[1:]:
            row_dict = {}
            for h, v in zip(headers, row):
                val = str(v).strip() if v is not None else ""
                row_dict[h] = val
            data.append(row_dict)
            
        self.excel_data = data
        return self.excel_data
    
    def load_docx(self) -> List[Dict[str, str]]:
        """Load DOCX file and extract text chunks."""
        doc = Document(self.docx_path)
        chunks = []
        
        # Extract paragraphs
        current_section = ""
        paragraph_buffer = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading (simple heuristic)
            if len(text) < 100 and (text.isupper() or text.endswith(':')):
                # Save previous section
                if paragraph_buffer:
                    chunks.append({
                        'type': 'paragraph',
                        'section': current_section,
                        'content': ' '.join(paragraph_buffer),
                        'source': 'Pharma_Clinical_Trial_Notes.docx'
                    })
                    paragraph_buffer = []
                current_section = text
            else:
                paragraph_buffer.append(text)
        
        # Add final buffer
        if paragraph_buffer:
            chunks.append({
                'type': 'paragraph',
                'section': current_section,
                'content': ' '.join(paragraph_buffer),
                'source': 'Pharma_Clinical_Trial_Notes.docx'
            })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            # Get headers from first row
            if not table.rows:
                continue
                
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            # Process each row
            for row_idx, row in enumerate(table.rows[1:], start=1):
                row_data = {}
                row_text_parts = []
                
                for header, cell in zip(headers, row.cells):
                    cell_text = cell.text.strip()
                    row_data[header] = cell_text
                    if cell_text:
                        row_text_parts.append(f"{header}: {cell_text}")
                
                if row_text_parts:
                    chunks.append({
                        'type': 'table',
                        'table_index': table_idx,
                        'row_index': row_idx,
                        'content': ' | '.join(row_text_parts),
                        'structured_data': row_data,
                        'source': 'Pharma_Clinical_Trial_Notes.docx'
                    })
        
        self.doc_chunks = chunks
        return chunks
    
    def get_excel_chunks(self) -> List[Dict[str, str]]:
        """Convert Excel rows to searchable text chunks."""
        if not self.excel_data:
            self.load_excel()
        
        chunks = []
        for idx, row in enumerate(self.excel_data):
            # Create a text representation of the row
            text_parts = []
            for col, value in row.items():
                if value and str(value).strip():
                    text_parts.append(f"{col}: {value}")
            
            chunks.append({
                'type': 'excel_row',
                'row_index': idx,
                'content': ' | '.join(text_parts),
                'structured_data': row,
                'source': 'Pharma_Clinical_Trial_AllDrugs.xlsx'
            })
        
        return chunks
    
    def load_all(self) -> Tuple[List[Dict], List[Dict], List[Dict]]:
        """Load all sources and return organized data."""
        excel_data = self.load_excel()
        doc_chunks = self.load_docx()
        excel_chunks = self.get_excel_chunks()
        
        return excel_data, doc_chunks, excel_chunks
